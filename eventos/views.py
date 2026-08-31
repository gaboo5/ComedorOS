from django.shortcuts import render
from django.db.models import Sum
from .models import Evento, Periodo


def agenda(request):
    periodo_id = request.GET.get('periodo')
    periodos = Periodo.objects.all()
    eventos = Evento.objects.select_related('tipo', 'periodo').order_by('fecha')
    if periodo_id:
        eventos = eventos.filter(periodo_id=periodo_id)

    context = {
        'eventos': eventos,
        'periodos': periodos,
        'periodo_seleccionado': periodo_id,
    }
    return render(request, 'eventos/agenda.html', context)


def dashboard(request):
    periodo_id = request.GET.get('periodo')
    periodos = Periodo.objects.all()
    eventos = Evento.objects.all()
    if periodo_id:
        eventos = eventos.filter(periodo_id=periodo_id)

    internos = eventos.filter(origen='interno')
    externos = eventos.filter(origen='externo')

    costo_interno = internos.aggregate(t=Sum('costo'))['t'] or 0
    costo_externo = externos.aggregate(t=Sum('costo'))['t'] or 0
    facturado_interno = internos.aggregate(t=Sum('precio_facturado'))['t'] or 0
    facturado_externo = externos.aggregate(t=Sum('precio_facturado'))['t'] or 0

    context = {
        'periodos': periodos,
        'periodo_seleccionado': periodo_id,
        'total_eventos': eventos.count(),
        'total_internos': internos.count(),
        'total_externos': externos.count(),
        'costo_interno': costo_interno,
        'costo_externo': costo_externo,
        'margen_interno': facturado_interno - costo_interno,
        'margen_externo': facturado_externo - costo_externo,
        'cobrados': eventos.filter(estado_facturacion='cobrado').count(),
        'pendientes': eventos.filter(estado_facturacion='pendiente').count(),
        'balance': (facturado_interno + facturado_externo) - (costo_interno + costo_externo),
    }
    return render(request, 'eventos/dashboard.html', context)

from django.shortcuts import redirect, get_object_or_404
from .forms import EventoForm


def evento_crear(request):
    if request.method == 'POST':
        form = EventoForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('agenda')
    else:
        form = EventoForm()
    return render(request, 'eventos/evento_form.html', {'form': form, 'titulo': 'Nuevo evento'})


def evento_editar(request, pk):
    evento = get_object_or_404(Evento, pk=pk)
    if request.method == 'POST':
        form = EventoForm(request.POST, instance=evento)
        if form.is_valid():
            form.save()
            return redirect('agenda')
    else:
        form = EventoForm(instance=evento)
    return render(request, 'eventos/evento_form.html', {'form': form, 'titulo': 'Editar evento'})


def evento_eliminar(request, pk):
    evento = get_object_or_404(Evento, pk=pk)
    if request.method == 'POST':
        evento.delete()
        return redirect('agenda')
    return render(request, 'eventos/evento_confirmar_eliminar.html', {'evento': evento})

from django.db import transaction
from .models import Presupuesto, Periodo
from .forms import EventoPresupuestoForm, PresupuestoForm, OpcionMenuFormSet


def presupuesto_form(request, pk=None):
    if pk:
        presupuesto = get_object_or_404(Presupuesto, pk=pk)
        evento = presupuesto.evento
    else:
        presupuesto = Presupuesto()
        evento = Evento()

    if request.method == 'POST':
        evento_form = EventoPresupuestoForm(request.POST, instance=evento)
        presupuesto_form_obj = PresupuestoForm(request.POST, instance=presupuesto)
        formset = OpcionMenuFormSet(request.POST, instance=presupuesto)

        if evento_form.is_valid() and presupuesto_form_obj.is_valid() and formset.is_valid():
            with transaction.atomic():
                evento = evento_form.save(commit=False)
                semestre = 1 if evento.fecha.month <= 6 else 2
                periodo, _ = Periodo.objects.get_or_create(anio=evento.fecha.year, semestre=semestre)
                evento.periodo = periodo
                nombre_partes = [str(evento.tipo_servicio)] if evento.tipo_servicio else []
                if evento.cliente:
                    nombre_partes.append(evento.cliente)
                evento.nombre = ' - '.join(nombre_partes) or 'Evento sin nombre'
                if not pk:
                    evento.estado_evento = 'presupuestado'
                evento.save()

                presupuesto_obj = presupuesto_form_obj.save(commit=False)
                presupuesto_obj.evento = evento
                presupuesto_obj.save()
                presupuesto_form_obj.save_m2m()

                formset.instance = presupuesto_obj
                formset.save()

            return redirect('presupuesto_detalle', pk=presupuesto_obj.pk)
    else:
        evento_form = EventoPresupuestoForm(instance=evento)
        presupuesto_form_obj = PresupuestoForm(instance=presupuesto)
        formset = OpcionMenuFormSet(instance=presupuesto)

    return render(request, 'eventos/presupuesto_form.html', {
        'evento_form': evento_form,
        'presupuesto_form': presupuesto_form_obj,
        'formset': formset,
        'editando': pk is not None,
    })


def presupuesto_lista(request):
    presupuestos = Presupuesto.objects.select_related('evento').order_by('-fecha_creacion')
    return render(request, 'eventos/presupuesto_lista.html', {'presupuestos': presupuestos})


def presupuesto_detalle(request, pk):
    presupuesto = get_object_or_404(Presupuesto, pk=pk)
    return render(request, 'eventos/presupuesto_detalle.html', {'presupuesto': presupuesto})

from io import BytesIO
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def presupuesto_exportar_word(request, pk):
    presupuesto = get_object_or_404(Presupuesto, pk=pk)
    evento = presupuesto.evento

    doc = Document()

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('PRESUPUESTO')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # Datos generales
    p = doc.add_paragraph()
    p.add_run('SOLICITANTE: ').bold = True
    p.add_run(evento.solicitante or '-')

    p = doc.add_paragraph()
    p.add_run('CLIENTE / ORGANISMO: ').bold = True
    p.add_run(evento.cliente or '-')

    p = doc.add_paragraph()
    p.add_run('FECHA: ').bold = True
    p.add_run(evento.fecha.strftime('%d de %B de %Y'))

    p = doc.add_paragraph()
    p.add_run('LUGAR: ').bold = True
    p.add_run(presupuesto.lugar or '-')

    if evento.cantidad_personas:
        p = doc.add_paragraph()
        p.add_run('CANTIDAD DE PERSONAS: ').bold = True
        texto_personas = str(evento.cantidad_personas)
        if presupuesto.cantidad_personas_celiaco:
            texto_personas += f' ({presupuesto.cantidad_personas_celiaco} celíacos)'
        p.add_run(texto_personas)

    doc.add_paragraph()

    # Opciones de menú
    for opcion in presupuesto.opciones.all():
        p = doc.add_paragraph()
        run = p.add_run(f'OPCIÓN {opcion.orden}: ' if opcion.orden else f'{opcion.nombre.upper()}: ')
        run.bold = True
        run.italic = True

        for linea in opcion.detalle_menu.splitlines():
            linea = linea.strip()
            if linea:
                p_plato = doc.add_paragraph()
                p_plato.add_run(linea).italic = True

        doc.add_paragraph()

    # Servicios incluidos
    p = doc.add_paragraph()
    p.add_run('Además, el servicio incluye:').bold = True

    extras = list(presupuesto.extras_incluidos.all())
    if extras:
        doc.add_paragraph(', '.join(str(e) for e in extras))
    if presupuesto.otros_extras:
        doc.add_paragraph(presupuesto.otros_extras)

    doc.add_paragraph()

    # Precios por opción
    p = doc.add_paragraph()
    p.add_run('PRECIO POR PERSONA:').bold = True

    for opcion in presupuesto.opciones.all():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        nombre_op = f'Opción {opcion.orden}' if opcion.orden else opcion.nombre
        run = p.add_run(f'{nombre_op}: ')
        run.bold = True
        p.add_run(f'${opcion.precio_por_persona:,.2f} por persona')

    doc.add_paragraph()

    # Notas
    if presupuesto.recargo_celiaco_por_persona:
        p = doc.add_paragraph()
        run = p.add_run(
            f'Nota: Las opciones de menú para celíacos tendrán un valor de '
            f'${presupuesto.recargo_celiaco_por_persona:,.2f} adicional al valor de la opción elegida.'
        )
        run.bold = True

    p = doc.add_paragraph()
    run = p.add_run(
        'Nota: El presente presupuesto puede sufrir modificaciones al momento de realizarse el evento. '
        'Se actualizará cada 30 días o cuando sea necesario.'
    )
    run.bold = True

    if presupuesto.observaciones:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Observaciones: ').bold = True
        p.add_run(presupuesto.observaciones)

    # Guardar en memoria y devolver como descarga
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre_archivo = f'Presupuesto_{evento.cliente or "evento"}_{evento.fecha}.docx'.replace(' ', '_')
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
    return response

from django.db.models import Sum, Count
from .models import InformeGestion
from .forms import InformeGestionForm


def calcular_indicadores_periodo(periodo):
    eventos = Evento.objects.filter(periodo=periodo)

    totales = eventos.aggregate(
        total_costo=Sum('costo'),
        total_facturado=Sum('precio_facturado'),
        total_personas=Sum('cantidad_personas'),
    )
    costo = totales['total_costo'] or 0
    facturado = totales['total_facturado'] or 0

    servicios_por_tipo = list(
        eventos.exclude(tipo_servicio__isnull=True)
        .values('tipo_servicio__nombre')
        .annotate(cantidad=Count('id'))
        .order_by('-cantidad')
    )

    unidades_destino = eventos.exclude(cliente='').values('cliente').distinct().count()

    return {
        'periodo': periodo,
        'total_eventos': eventos.count(),
        'internos': eventos.filter(origen='interno').count(),
        'externos': eventos.filter(origen='externo').count(),
        'total_personas': totales['total_personas'] or 0,
        'costo_total': costo,
        'facturado_total': facturado,
        'margen_total': facturado - costo,
        'cobrados': eventos.filter(estado_facturacion='cobrado').count(),
        'pendientes': eventos.filter(estado_facturacion='pendiente').count(),
        'servicios_por_tipo': servicios_por_tipo,
        'unidades_destino': unidades_destino,
    }


def informe_gestion_detalle(request, periodo_id):
    periodo = get_object_or_404(Periodo, pk=periodo_id)
    indicadores = calcular_indicadores_periodo(periodo)
    informe, _ = InformeGestion.objects.get_or_create(periodo=periodo)

    return render(request, 'eventos/informe_gestion_detalle.html', {
        'indicadores': indicadores,
        'informe': informe,
    })


def informe_gestion_form(request, periodo_id):
    periodo = get_object_or_404(Periodo, pk=periodo_id)
    informe, _ = InformeGestion.objects.get_or_create(periodo=periodo)

    if request.method == 'POST':
        form = InformeGestionForm(request.POST, instance=informe)
        if form.is_valid():
            form.save()
            return redirect('informe_gestion_detalle', periodo_id=periodo.id)
    else:
        form = InformeGestionForm(instance=informe)

    return render(request, 'eventos/informe_gestion_form.html', {'form': form, 'periodo': periodo})

def informe_gestion_exportar_word(request, periodo_id):
    periodo = get_object_or_404(Periodo, pk=periodo_id)
    indicadores = calcular_indicadores_periodo(periodo)
    informe, _ = InformeGestion.objects.get_or_create(periodo=periodo)

    doc = Document()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f'INFORME DE GESTIÓN — {periodo}')
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()

    def agregar_titulo(texto):
        p = doc.add_paragraph()
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(13)

    def agregar_texto(texto):
        doc.add_paragraph(texto or '(sin completar)')

    agregar_titulo('1. Resumen Ejecutivo')
    agregar_texto(informe.resumen_ejecutivo)

    agregar_titulo('2. Objetivos del Período')
    agregar_texto(informe.objetivos)

    agregar_titulo('3. Indicadores Clave y Cuadro Operativo')
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = 'Light Grid Accent 1'
    hdr = tabla.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Categoría / Métrica', 'Detalle / Unidad', 'Total del Período'

    filas = [
        ('Eventos Atendidos', 'Internos / Externos',
         f"{indicadores['total_eventos']} eventos ({indicadores['internos']} internos, {indicadores['externos']} externos)"),
        ('Raciones / Menús Especiales', 'Común, vegetariano, celíaco',
         f"{informe.raciones_comun or 0} común, {informe.raciones_vegetariano or 0} vegetariano, {informe.raciones_celiaco or 0} celíaco"),
        ('Asistencia Estimada', 'Personas atendidas', f"{indicadores['total_personas']} personas"),
        ('Servicios Adicionales', 'Horas extra / servicios', str(informe.servicios_adicionales_horas_extra or 0)),
        ('Unidades de Destino', 'Facultades, dependencias, Rectorado', str(indicadores['unidades_destino'])),
    ]
    for cat, detalle, total in filas:
        fila = tabla.add_row().cells
        fila[0].text, fila[1].text, fila[2].text = cat, detalle, total

    doc.add_paragraph()
    agregar_titulo('4. Desglose de Servicios Realizados')
    if indicadores['servicios_por_tipo']:
        for item in indicadores['servicios_por_tipo']:
            doc.add_paragraph(f"{item['tipo_servicio__nombre']}: {item['cantidad']} eventos", style='List Bullet')
    else:
        doc.add_paragraph('(sin datos de tipo de servicio para este período)')

    agregar_titulo('5. Balance Económico y Financiero')
    doc.add_paragraph(f"Presupuesto asignado: ${informe.presupuesto_asignado or 0:,.2f}")
    doc.add_paragraph(f"Ejecución presupuestaria (costos operativos): ${indicadores['costo_total']:,.2f}")
    doc.add_paragraph(f"Facturado: ${indicadores['facturado_total']:,.2f} "
                       f"({indicadores['cobrados']} cobrados, {indicadores['pendientes']} pendientes)")
    doc.add_paragraph(f"Margen / Balance: ${indicadores['margen_total']:,.2f}")
    if informe.presupuesto_asignado:
        porcentaje = (indicadores['costo_total'] / informe.presupuesto_asignado) * 100
        doc.add_paragraph(f"Porcentaje de ejecución del presupuesto asignado: {porcentaje:.1f}%")

    agregar_titulo('6. Logros Destacados y Mejoras Operativas')
    agregar_texto(informe.logros_destacados)

    agregar_titulo('7. Desafíos y Objetivos para el Próximo Período')
    agregar_texto(informe.desafios_proximo_periodo)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre_archivo = f'Informe_Gestion_{periodo.anio}_S{periodo.semestre}.docx'
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
    return response